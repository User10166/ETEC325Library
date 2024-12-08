import numpy as np
import pandas as pd
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

stdoutInstance = sys.stdout

class WordDocument():
    def __init__(self, file):
        self.file = file + ".docx";
        self.document = Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        style = self.document.styles['No Spacing']
        font = style.font
        font.name = 'JetBrains Mono'
        font.size = Pt(13)
        self.text = ''

    def writeText(self, text):
        if(self.text == ''):
            self.text = text
        else:
            self.text = self.text + text

    def saveDocument(self):
        paragraph = self.document.add_paragraph(self.text)
        paragraph.style = self.document.styles['No Spacing']
        self.document.save(self.file)

class Logger(object):
    def __init__(self, file):
        self.terminal = sys.stdout
        self.wordDocument = WordDocument(file)
   
    def write(self, message):
        self.terminal.write(message)
        self.wordDocument.writeText(message)

    def closeLog(self):
        self.wordDocument.saveDocument()

    def flush(self):
        # this flush method is needed for python 3 compatibility.
        # this handles the flush command by doing nothing.
        # you might want to specify some extra behavior here.
        pass

    def __del__(self):
        self.closeLog()
        sys.stdout = stdoutInstance

def enableLogging(file):
    sys.stdout = Logger(file)

def disableLogging():
    sys.stdout.closeLog()
    sys.stdout = stdoutInstance

def createRange(start, stop, step=1):
    inclusiveRange = np.arange(start, stop + step, step)
    rangeShape = inclusiveRange.shape
    return np.reshape(inclusiveRange, (1, rangeShape[0]))
    
def printTable(array, columnNames=None):
    arrayShape = array.shape
    if(len(arrayShape) == 1):
        array = np.reshape(array, (1, arrayShape[0]))
        arrayShape = array.shape
    if(columnNames is None):
        if(len(arrayShape) == 1):
            numRows = int(arrayShape[0])
            blankRows = ['' for i in range(numRows)]
            df = pd.DataFrame(array, index=blankRows, columns=[''])
        elif(len(arrayShape) > 1):
            numRows = int(arrayShape[0])
            blankRows = ['' for i in range(numRows)]
            numColumns = int(arrayShape[1])
            blankColumns = ['' for i in range(numColumns)]
            df = pd.DataFrame(array, index=blankRows, columns=blankColumns)
    else:
        numRows = int(arrayShape[0])
        blankRows = ['' for i in range(numRows)]
        df = pd.DataFrame(array, index=blankRows, columns=columnNames)
    pd.set_option("display.max_rows", None)
    pd.set_option("display.max_columns", None)
    print(df)
    
def printVariable(label, var, columnNames=None):
    if(type(var) is np.ndarray):
        print("\n" + label + ":")
        printTable(var, columnNames)
    else:
        print("\n" + label + ": " + str(var))
        
